from docx.enum.text import WD_ALIGN_PARAGRAPH


def find_parageaphs_sn(doc,need_find_text):  #返回需要查找的文字段落
	for par_sn in range(0,len(doc.paragraphs)):
		if need_find_text in doc.paragraphs[par_sn].text:
			return par_sn


def fix_table(doc,sn_table,table_row,table_col,add_text,style):
	table = doc.tables[sn_table]
	table.style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

	cell = table.cell(table_row,table_col).paragraphs[0].add_run(text=add_text)
	cell.font.bold = style['bold']
	cell.font.size = style['size']
	cell.font.name = style['name']


def set_font(text, text_bold, text_size, text_name):
	text.font.bold = text_bold  # 对文本对象加粗
	text.font.size = text_size  # 设置文本文字大小
	text.font.name = text_name  # 设置文本文字字体

def table_add_rows(doc,sn_tables):
	doc.tables[sn_tables].add_row()
