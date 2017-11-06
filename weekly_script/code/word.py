import datetime
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from def_files.word_def import  *
from docx.shared import Mm,Pt
from get_data import get_from_xlsx


#doc = docx.Document('C:\\Users\\Administrator\\Desktop\\周报模版.docx')
doc = docx.Document('F:\github\python\weekly_script\code\周报模版.docx')
end_time = datetime.datetime.now() + datetime.timedelta(days=-1)#结束的那一天
begin_time = end_time + datetime.timedelta(days=-7)#开始的那一天

####################修改“X年X月X日至X年X月X日”############################
first_fix = '维护工作周报'#第一个需要修改的地方
first_fix_par_sn =find_parageaphs_sn(doc,first_fix) + 1#第一个要修改的地方行号
weekly_time = '%s年%s月%s日至%s年%s月%s日'%(begin_time.year,
	begin_time.month,
	begin_time.day,
	end_time.year,
	end_time.month,
	end_time.day)

date_par = doc.paragraphs[first_fix_par_sn]  #返回行对象
date_par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER#居中

date_text = date_par.add_run(text=weekly_time)#返回行里面的文本对象
set_font(date_text,True,Mm(3.8805555555555555),'宋体')

##########################修改“X月第X周”##################################
two_fix_par_sn = first_fix_par_sn + 7  #在第一次修改的行数往下数7行进行修改
what_week = '%s月份第%s周'%(begin_time.month,
	int(datetime.datetime(
		end_time.year,end_time.month,end_time.day).strftime("%W")
	)
	-
	int(datetime.datetime(
		begin_time.year,begin_time.month,1).strftime("%W")
	))

two_fix_par = doc.paragraphs[two_fix_par_sn]  #返回行对象
two_fix_par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER#进行居中

two_date = two_fix_par.add_run(text=what_week) #返回行里面的文本对象
set_font(two_date,True,Mm(12.70),'宋体')

########################修改编制时间:X年X月X日############################
three_fix = '编制时间'
three_par_sn = find_parageaphs_sn(doc,three_fix)

three_fix_par = doc.paragraphs[three_par_sn]#返回行对象
three_add_date = three_fix_par.add_run(
	text='%s年%s月%s日'%(end_time.year,end_time.month,end_time.day))

set_font(three_add_date,False,Mm(4.94),'宋体')

#########################修改文档资料信息表格##############################
cell_style = {'bold':False,'size':Pt(10.5),'name':'宋体'}

document_name = '虚拟化平台周报_%s%s%s'%(end_time.year,end_time.month,end_time.day)
fix_table(doc,0,0,1,document_name,cell_style)

document_change_time = '%s年%s月%s日'%(end_time.year,end_time.month,end_time.day)
fix_table(doc,0,3,1,document_change_time,cell_style)

########################修改检查工作安排表格##################################
cell_style = {'bold':False,'size':Pt(10.5),'name':'宋体'}

inspect_begin_time = '%s年%s月%s日'%(begin_time.year,begin_time.month,begin_time.day)
inspect_end_time = '%s年%s月%s日'%(end_time.year,end_time.month,end_time.day)

fix_table(doc,1,0,2,inspect_begin_time,cell_style)
fix_table(doc,1,1,2,inspect_end_time,cell_style)

########################修改工单统计表格##################################
event_tol = get_from_xlsx.event

cell_style = {'bold':False,'size':Pt(10.5),'name':'宋体'}

number = 1
rows = 1

for event in event_tol:
	table_add_rows(doc,4)
	fix_table(doc,4,rows,0,str(number),cell_style)
	fix_table(doc,4,rows,1,event,cell_style)
	fix_table(doc,4,rows,2,event_tol[event][0].split('/')[0],cell_style)#类别
	fix_table(doc,4,rows,3,event_tol[event][1],cell_style)
	fix_table(doc,4,rows,4,'已完成',cell_style)
	rows = rows + 1
	number = number + 1

########################修改虚拟机表格####################################
event_tol = get_from_xlsx.event
cell_style = {'bold':False,'size':Pt(10.5),'name':'宋体'}

number = 1
rows = 1
vms = 0   #用于后面的问题统计中的虚拟机数量

for event in event_tol:
	table_add_rows(doc,6)#添加一行
	vms = vms + 1
	fix_table(doc,6,rows,0,str(number),cell_style)#序号
	fix_table(doc,6,rows,1,event_tol[event][0].split('/')[2],cell_style)#平台名
	fix_table(doc,6,rows,2,event_tol[event][0].split('/')[1],cell_style)#故障范围
	fix_table(doc,6,rows,3,event_tol[event][1],cell_style)#问题描述
	fix_table(doc,6,rows,4,'已完成', cell_style)
	fix_table(doc,6,rows,5,'增强巡检，增加监控',cell_style)
	rows= rows + 1
	number = number + 1

########################修改存储表格####################################
datastore_tol = get_from_xlsx.datastore_tol

row_id = 1
cell_style = {'bold':False,'size':Pt(10.5),'name':'宋体'}

for row in datastore_tol:
	table_add_rows(doc,7) #添加一行
	fix_table(doc,7,row_id,0,row['sn'],cell_style)  #序号
	fix_table(doc,7,row_id,1,row['datastore_name'],cell_style)#存储名
	fix_table(doc,7,row_id,2,row['datastore_type'],cell_style)#存储类型
	fix_table(doc,7,row_id,3,row['problem_description'],cell_style)#问题描述
	fix_table(doc,7,row_id,4,row['free'], cell_style)    #存储空闲率
	fix_table(doc,7,row_id,5,row['problem_state'],cell_style)#问题状态
	fix_table(doc,7,row_id,6,row['problem_result'],cell_style)#处理结果
	row_id = row_id + 1  #行号加1

#############################最后修改问题统计##################################
cell_style = {'bold':False,'size':Pt(10.5),'name':'宋体'}

fix_table(doc,3,4,3,str(vms)+'台',cell_style)

########################保存word文件####################################

weekly_filename = "C:\\Users\\Administrator\\Desktop\\周报%s.docx"%(str(end_time.year)+str(end_time.month)+str(end_time.day))

doc.save(weekly_filename)
